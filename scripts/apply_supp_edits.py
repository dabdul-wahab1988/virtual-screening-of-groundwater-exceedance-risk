# -*- coding: utf-8 -*-
"""Add Supplementary Tables S8 (bootstrap CIs) and S9 (spatial k-sensitivity),
and extend Supplementary Methods S7 with clustering-algorithm detail."""
import docx
from docx.shared import Pt
import pandas as pd

SRC = "../supplementary8062026.docx"
DST = "../supplementary8062026_revised.docx"

d = docx.Document(SRC)
paras = d.paragraphs


def append_sentence(para_idx, text):
    p = paras[para_idx]
    run = p.add_run(" " + text)
    if p.runs:
        ref = p.runs[0]
        run.font.size = ref.font.size
        run.font.name = ref.font.name
    else:
        run.font.size = Pt(10)


# Extend Methods S7 (cross-validation designs) with clustering-algorithm detail
append_sentence(31,
    "Spatial clusters were constructed using k-means partitioning on great-circle (Haversine) "
    "distances between well coordinates, with cluster centroids updated as the coordinate-space "
    "mean of member wells (a negligible approximation given the <100 km extent of the study area). "
    "The number of clusters (k) was selected automatically via an inertia-elbow heuristic, stopping "
    "once the marginal reduction in within-cluster sum of squared Haversine distances fell below "
    "10% of the total range, capped at eight clusters; this yielded k = 4 for the Pru Basin dataset. "
    "A sensitivity analysis re-fitting spatial-group cross-validation at k = 3, 5, 6, 7 and 8 is "
    "reported in Supplementary Table S9.")

# ---------------------------------------------------------------------
# Locate anchor: "Supplementary Figures" heading, to insert new content before it
anchor = None
for p in paras:
    if p.text.strip() == "Supplementary Figures":
        anchor = p
        break
assert anchor is not None, "Could not find 'Supplementary Figures' heading"
anchor_elem = anchor._p

heading_style = None
for p in paras:
    if p.text.strip().startswith("Supplementary Table S7"):
        heading_style = p.style
        break
body_style = paras[94].style  # S7 footnote paragraph style, reused for new footnotes


def make_heading(text):
    np_ = d.add_paragraph(text)
    if heading_style is not None:
        np_.style = heading_style
    anchor_elem.addprevious(np_._p)
    return np_


def make_body(text):
    np_ = d.add_paragraph(text)
    np_.style = body_style
    anchor_elem.addprevious(np_._p)
    return np_


def make_table(nrows, ncols):
    t = d.add_table(rows=nrows, cols=ncols)
    t.style = d.tables[22].style
    anchor_elem.addprevious(t._tbl)
    return t


def fill_table(t, header, rows):
    for j, h in enumerate(header):
        t.rows[0].cells[j].text = h
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = str(val)
    for r in t.rows:
        for c in r.cells:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)


# --- Supplementary Table S8: bootstrap CIs --------------------------------
make_heading("Supplementary Table S8 | Bootstrap 95% confidence intervals for Table 3 headline classifiers.")

ci = pd.read_csv("outputs/r_exports/bootstrap_ci_summary.csv")
tier_label = {
    "Tier1_Field": "Tier 1: field",
    "Tier2_Reduced": "Tier 2: reduced",
    "Tier3_Full": "Tier 3: full",
    "Tier2_Reduced_TDS_EC_strict": "Tier 2: EC-strict",
}
algo_label = {
    "EcOnlyLogistic": "EC-only logistic",
    "LogisticRegression": "Logistic regression",
    "RandomForest": "Random forest",
    "GradientBoostedTrees": "Gradient-boosted trees",
}
rows = []
for _, r in ci.iterrows():
    rows.append([
        r["target"], tier_label.get(r["tier"], r["tier"]), algo_label.get(r["algorithm"], r["algorithm"]),
        f"{r['pr_auc_median']:.3f} [{r['pr_auc_ci_lo']:.3f}, {r['pr_auc_ci_hi']:.3f}]",
        f"{r['recall_median']:.3f} [{r['recall_ci_lo']:.3f}, {r['recall_ci_hi']:.3f}]",
        f"{r['f2_median']:.3f} [{r['f2_ci_lo']:.3f}, {r['f2_ci_hi']:.3f}]",
        f"{r['brier_median']:.4f} [{r['brier_ci_lo']:.4f}, {r['brier_ci_hi']:.4f}]",
    ])

t8 = make_table(len(rows) + 1, 7)
fill_table(t8, ["Target", "Predictor tier", "Algorithm", "PR-AUC (95% CI)", "Recall (95% CI)",
                "F2 (95% CI)", "Brier (95% CI)"], rows)

make_body(
    "Values are bootstrap medians with 95% percentile confidence intervals (2000 resamples), "
    "obtained by resampling wells with replacement from the stratified nested cross-validation "
    "out-of-fold predictions underlying Table 3 and, for each resampled well, drawing one of its "
    "repeat-fold predictions at random. Bootstrap medians are computed on resampled prediction sets "
    "and therefore differ slightly from the single aggregate point estimates reported in Table 3; "
    "the two are complementary, with Table 3 giving the point estimate on the full out-of-fold "
    "prediction set and this table giving the associated resampling uncertainty. The TDS EC-strict "
    "row corresponds to the additional sensitivity classifier reported in Table 3 that excludes EC "
    "as a conductivity proxy. PR-AUC, precision-recall area under the curve; F2, recall-weighted "
    "F-score."
)

# --- Supplementary Table S9: spatial cluster sensitivity ------------------
make_heading("Supplementary Table S9 | Sensitivity of spatial-group cross-validation PR-AUC to the number of spatial clusters (k).")

ks = pd.read_csv("outputs/r_exports/spatial_k_sensitivity_summary.csv")
rows9 = []
for _, r in ks.iterrows():
    def fmt(v):
        return "NE" if pd.isna(v) else f"{v:.3f}"
    rows9.append([
        r["target"], tier_label.get(r["tier"], r["tier"]), algo_label.get(r["algorithm"], r["algorithm"]),
        fmt(r["3"]), fmt(r["4"]), fmt(r["5"]), fmt(r["6"]), fmt(r["7"]), fmt(r["8"]),
        "NE" if pd.isna(r["range"]) else f"{r['range']:.3f}",
    ])

t9 = make_table(len(rows9) + 1, 10)
fill_table(t9, ["Target", "Predictor tier", "Algorithm", "k=3", "k=4 (main text)", "k=5", "k=6", "k=7",
                "k=8", "Range"], rows9)

make_body(
    "Values are mean spatial-group cross-validation PR-AUC for the same target/tier/algorithm "
    "combination reported as the Table 3 headline classifier, re-fitted after reconstructing "
    "spatial clusters at each alternative k using the same k-means/Haversine procedure described "
    "in Supplementary Methods S7. k=4 is the elbow-selected value used throughout the main text. "
    "NE indicates the combination was not evaluable at that k: k=2 could not support the three "
    "held-out spatial folds required by the validation design, and TDS at k=7 produced one held-out "
    "spatial group with zero TDS exceedances. Range is the maximum minus minimum PR-AUC across all "
    "evaluable k for that target/tier/algorithm. Na, Cl and TDS show narrow ranges, indicating that "
    "their spatial-validation performance is not an artefact of the specific k=4 partition; B and F "
    "show wider ranges, consistent with their greater spatial sensitivity reported in the main text."
)

d.save(DST)
print("Saved:", DST)
