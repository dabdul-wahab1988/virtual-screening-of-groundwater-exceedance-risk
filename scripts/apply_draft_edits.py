# -*- coding: utf-8 -*-
"""Apply reviewer-response edits to draft8062026.docx, saved as a new revised file."""
import docx
from docx.shared import Pt
import copy

SRC = "../draft8062026.docx"
DST = "../draft8062026_revised.docx"

d = docx.Document(SRC)
paras = d.paragraphs


def append_sentence(para_idx, text):
    p = paras[para_idx]
    # clone formatting from the paragraph's last run
    run = p.add_run(" " + text)
    if p.runs:
        ref = p.runs[0]
        run.font.size = ref.font.size
        run.font.name = ref.font.name
        run.italic = ref.font.italic
    else:
        run.font.size = Pt(10)


# --- Q1: foreground EC-strict TDS sensitivity result ---------------------
append_sentence(22,  # 2.3 Predictor-tier design
    "The EC-strict TDS result is treated as the primary evidence against EC-proxy circularity "
    "and is reported as a distinct row in Table 3 alongside the EC-inclusive, field-screening estimate.")

append_sentence(52,  # 3.4 Model performance, TDS/Cl paragraph
    "Because EC is itself a field-scale proxy for total dissolved ionic content, the EC-only PR-AUC "
    "of 1.000 for TDS could partly reflect proxy circularity; under the EC-strict sensitivity design "
    "(EC excluded), the best non-EC classifier (gradient-boosted trees) still achieved a PR-AUC of "
    "0.975 (95% bootstrap CI 0.879–1.000) under stratified nested cross-validation and a spatial "
    "PR-AUC of 0.942, confirming that TDS screenability is not solely an artefact of the EC–TDS "
    "relationship (Table 3).")

append_sentence(73,  # 4.2 Why salinity-related targets were more screenable
    "For TDS specifically, this strong EC-only performance should be read alongside the EC-strict "
    "sensitivity result, which shows that most of TDS screenability persists even when the EC proxy "
    "is removed (Table 3).")

# --- Q2: spatial clustering methodology + k-sensitivity -------------------
append_sentence(28,  # 2.6 Spatial grouping
    "Clusters were constructed using k-means partitioning on great-circle (Haversine) distances "
    "between well coordinates, with the number of clusters selected automatically via an "
    "inertia-elbow rule (stopping once the marginal reduction in within-cluster sum of squared "
    "distances fell below 10% of the total range, capped at eight clusters), yielding four clusters "
    "for this dataset. To test the sensitivity of spatial-validation performance to this choice, "
    "models were re-fitted under spatial-group cross-validation for alternative cluster counts "
    "(k = 3, 5, 6, 7, 8; Supplementary Table S9). Salinity-linked targets (Na, Cl, TDS) retained "
    "spatial PR-AUC within a narrow range across all tested k (≤0.02–0.03 spread for Na and Cl; "
    "TDS remained at or near 1.000 wherever evaluable), whereas B and F showed materially wider "
    "spread (PR-AUC ranges of 0.14–0.27), confirming that fluoride and, to a lesser extent, boron "
    "screening are more sensitive to the spatial partition than the salinity-linked targets. TDS "
    "could not be evaluated at k = 7 because one held-out spatial group contained zero TDS "
    "exceedances, illustrating the fold-instability risk associated with small spatial clusters "
    "noted in the Limitations.")

# --- SHAP -> path-based approximation caveat ------------------------------
append_sentence(39,  # 2.10 Operational thresholds, explainability and reproducibility
    "Feature attributions were computed as path-based (Saabas-style) decompositions of the fitted "
    "tree ensembles, an efficient approximation to exact Shapley values that is appropriate for the "
    "shallow trees used here but is not identical to exact TreeSHAP; this distinction should be "
    "considered when interpreting the reported attributions as approximate rather than exact.")

# --- Q3, Q4, Q5: Limitations additions ------------------------------------
append_sentence(84,  # 4.6 Limitations
    "Operational probability cutoffs (Table 4) were selected from pooled out-of-fold predictions "
    "after fold-level predictions had been generated, rather than being re-selected within each "
    "outer training fold; this pooled selection can be optimistic for the reported "
    "sensitivity/specificity at the chosen cutoff, although it does not affect the fold-computed "
    "PR-AUC, ROC-AUC or recall-at-default-threshold estimates in Table 3, which are estimated "
    "independently of the cutoff choice. Ninety-five percent bootstrap confidence intervals for the "
    "Table 3 headline classifiers, obtained by well-level resampling of the out-of-fold predictions, "
    "are reported in Supplementary Table S8 and show that PR-AUC values above roughly 0.90 in "
    "Table 3 (Na, Cl, TDS and the EC-strict TDS variant) remain above 0.85–0.88 at the lower 95% "
    "bound, whereas the F and B intervals are wide enough to include PR-AUC values below 0.5 at the "
    "lower bound, underscoring that these targets should be interpreted as screening-informative "
    "rather than precisely estimated. Below-detection-limit values for F⁻ and NO3⁻ were substituted "
    "using the detection-limit-divided-by-√2 convention before modelling (Supplementary Methods S2); "
    "this is distinct from the fold-wise median imputation applied later to the rare missing numeric "
    "entries during preprocessing (Section 2.5). Censoring-aware alternatives such as "
    "regression-on-order-statistics or Tobit models were not tested. For F⁻, all five "
    "below-detection-limit values were recorded at the assay's lower bound of 0.23 mg L⁻¹, roughly "
    "five-fold below the 1.125 mg L⁻¹ screening threshold, so the choice of substitution method is "
    "very unlikely to have altered exceedance labels or class separation near the decision boundary; "
    "NO3⁻ was not modelled beyond descriptive reporting, so BDL treatment for NO3⁻ does not affect "
    "any classifier.")

# --- Table 3: add EC-strict TDS row ---------------------------------------
t3 = d.tables[5]
new_row = t3.add_row()
values = ["TDS (strict, EC excluded)", "Tier 2: reduced", "Gradient-boosted trees",
          "0.975", "0.985", "0.920", "0.919", "0.0307", "0.954", "0.942"]
ref_cell = t3.rows[1].cells[0]
ref_size = ref_cell.paragraphs[0].runs[0].font.size if ref_cell.paragraphs[0].runs else Pt(10)
for cell, val in zip(new_row.cells, values):
    cell.text = val
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.size = ref_size

append_sentence(115,  # Table 3 note paragraph
    "For TDS, the EC-strict sensitivity design (EC excluded as a conductivity proxy) is reported "
    "as an additional row using the best non-EC classifier; it confirms that TDS remains screenable "
    "independent of the EC–TDS proxy relationship. Ninety-five percent bootstrap confidence "
    "intervals for all headline classifiers in this table are reported in Supplementary Table S8, "
    "and sensitivity of the spatial-validation estimates to the number of spatial clusters is "
    "reported in Supplementary Table S9.")

# --- Data and Code Availability statement ---------------------------------
# Insert before the References heading (paragraph index 91)
ref_heading = paras[91]
new_heading = ref_heading.insert_paragraph_before("Data and Code Availability")
new_heading.style = ref_heading.style if ref_heading.style.name.lower().startswith("heading") else new_heading.style
# Try to match a heading style used elsewhere (paragraph 87, "5. Conclusions")
try:
    new_heading.style = paras[87].style
except Exception:
    pass

body_text = (
    "The raw groundwater dataset originates from Sunkari et al. (2021). The leakage-control rules, "
    "predictor-tier definitions, spatial-cluster assignments, cross-validation fold assignments, "
    "hyperparameters, fold-level metrics, well-level predictions, SHAP/path-based attribution "
    "values, operational-threshold selections and GIS-ready exports referenced throughout this "
    "manuscript are stored in an auditable SQLite reproducibility database (Supplementary Methods "
    "S12). Code and fold-assignment/spatial-cluster metadata are available at "
    "[REPOSITORY URL / DOI — to be inserted by authors]."
)
new_body = ref_heading.insert_paragraph_before(body_text)
try:
    new_body.style = paras[88].style  # body-text style used in Conclusions
except Exception:
    pass

d.save(DST)
print("Saved:", DST)
